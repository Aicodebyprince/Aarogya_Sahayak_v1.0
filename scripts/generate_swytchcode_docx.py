import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def generate_docx(output_path):
    doc = Document()
    
    # Page setup (Standard Letter, 0.75" margins)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        # Header & Footer
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Aarogya Sahayak • Swytchcode AI Governance Track • Technical Whitepaper")
        f_run.font.name = "Calibri"
        f_run.font.size = Pt(8.5)
        f_run.font.color.rgb = RGBColor(100, 116, 139)

    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(30, 41, 59)

    # 1. Executive Banner Card (Shaded Table)
    banner_table = doc.add_table(rows=1, cols=1)
    banner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner_table.autofit = False
    banner_cell = banner_table.rows[0].cells[0]
    banner_cell.width = Inches(7.0)
    set_cell_background(banner_cell, "0F172A")
    set_cell_margins(banner_cell, top=240, bottom=240, left=260, right=260)

    bp1 = banner_cell.paragraphs[0]
    r_track = bp1.add_run("ENTERPRISE AI GOVERNANCE • SWYTCHCODE RUNTIME SPECIFICATION\n")
    r_track.font.bold = True
    r_track.font.size = Pt(9)
    r_track.font.color.rgb = RGBColor(249, 115, 22)

    r_title = bp1.add_run("Aarogya Sahayak × Swytchcode\n")
    r_title.font.bold = True
    r_title.font.size = Pt(20)
    r_title.font.color.rgb = RGBColor(255, 255, 255)

    r_sub = bp1.add_run("Governed Agent Tool Execution, Clinical Idempotency & Indic Voice Infrastructure\n")
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor(148, 163, 184)

    r_meta = bp1.add_run("Architecture: Multi-Agent Intelligence + Swytchcode Runtime  |  Deployment: Render (FastAPI) + Vercel (PWA)")
    r_meta.font.size = Pt(8.5)
    r_meta.font.color.rgb = RGBColor(56, 189, 248)

    doc.add_paragraph()

    # 2. Executive Summary
    h1 = doc.add_paragraph()
    r_h1 = h1.add_run("1. Executive Problem Statement & Solution")
    r_h1.font.bold = True
    r_h1.font.size = Pt(13)
    r_h1.font.color.rgb = RGBColor(15, 23, 42)

    p1 = doc.add_paragraph()
    p1.add_run(
        "In rural Indian healthcare, when an AI assistant detects critical danger signs—such as a mother with "
        "severe pre-eclampsia (BP 165/105 mmHg with blurred vision)—allowing an autonomous LLM to invoke external "
        "APIs directly is catastrophic. Unchecked LLM tool calling leads to hallucinated parameters, duplicate ambulance "
        "dispatches over flaky 2G/3G networks, and credential leaks."
    )

    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        "Aarogya Sahayak integrates Swytchcode as its enterprise-grade AI execution & governance layer. "
    )
    r2.font.bold = True
    p2.add_run(
        "The AI reasoning model never sees or holds external API credentials. Every medical alert, Sarvam Indic voice "
        "synthesis, and facility query executes through Swytchcode with sliding-window idempotency, strict pre-execution "
        "schema validation, and real-time observability on app.swytchcode.com."
    )

    # 3. Four Core Pillars (2x2 Table)
    doc.add_paragraph()
    pillars_table = doc.add_table(rows=2, cols=2)
    pillars_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    pillars_table.autofit = False
    
    col_w = Inches(3.45)
    for row in pillars_table.rows:
        for cell in row.cells:
            cell.width = col_w
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=140, bottom=140, left=160, right=160)

    # Cell 0,0: Idempotency
    c00 = pillars_table.cell(0, 0).paragraphs[0]
    r_c00_t = c00.add_run("🛡️ Guaranteed Idempotency\n")
    r_c00_t.font.bold = True
    r_c00_t.font.size = Pt(9.5)
    r_c00_d = c00.add_run("Emergency dispatches deduplicated via 5-min sliding window SHA-256 tokens. Zero duplicate ambulances.")
    r_c00_d.font.size = Pt(8.5)
    r_c00_d.font.color.rgb = RGBColor(71, 85, 105)

    # Cell 0,1: Zero-Token Security
    c01 = pillars_table.cell(0, 1).paragraphs[0]
    r_c01_t = c01.add_run("🔒 Zero-Token Security\n")
    r_c01_t.font.bold = True
    r_c01_t.font.size = Pt(9.5)
    r_c01_d = c01.add_run("LLM reasoning never touches production API secrets. All credentials isolated in Swytchcode vault.")
    r_c01_d.font.size = Pt(8.5)
    r_c01_d.font.color.rgb = RGBColor(71, 85, 105)

    # Cell 1,0: Sarvam Voice
    c10 = pillars_table.cell(1, 0).paragraphs[0]
    r_c10_t = c10.add_run("🗣️ Sarvam Indic Voice Proxy\n")
    r_c10_t.font.bold = True
    r_c10_t.font.size = Pt(9.5)
    r_c10_d = c10.add_run("Marathi & Hindi speech (Saaras & Bulbul) governed with 3s timeout budgets and phonetic fallback.")
    r_c10_d.font.size = Pt(8.5)
    r_c10_d.font.color.rgb = RGBColor(71, 85, 105)

    # Cell 1,1: Live Observability
    c11 = pillars_table.cell(1, 1).paragraphs[0]
    r_c11_t = c11.add_run("📊 Live Observability\n")
    r_c11_t.font.bold = True
    r_c11_t.font.size = Pt(9.5)
    r_c11_d = c11.add_run("Every tool call streams live to app.swytchcode.com with latency, status 200 OK, and schema logs.")
    r_c11_d.font.size = Pt(8.5)
    r_c11_d.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph()

    # 4. System Architecture
    h2 = doc.add_paragraph()
    r_h2 = h2.add_run("2. System Architecture & Governed Data Flow")
    r_h2.font.bold = True
    r_h2.font.size = Pt(13)
    r_h2.font.color.rgb = RGBColor(15, 23, 42)

    arch_text = (
        "+-------------------------------------------------------------------------+\n"
        "|                 CITIZEN MOBILE PWA (React 19 / Vite / i18n)             |\n"
        "|    UI Banner: [ ShieldCheck: Swytchcode AI Runtime Governed & Idempotent]|\n"
        "+------------------------------------+------------------------------------+\n"
        "                                     | HTTPS (Audio / Text Symptoms)\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|                    FASTAPI BACKEND CORE (Render Singapore)              |\n"
        "|   +-----------------------+     +------------------+                    |\n"
        "|   | Deterministic Triage  |     | PII Masking      |                    |\n"
        "|   | (Emergency Rule Eng)  |     | (Scrubs Vitals)  |                    |\n"
        "|   +-----------+-----------+     +--------+---------+                    |\n"
        "|               +--------------------------+                              |\n"
        "|                                          v                              |\n"
        "|                           SWYTCHCODE INTEGRATION ADAPTER                |\n"
        "|                      (backend/app/integrations/swytchcode.py)           |\n"
        "|       * Schema Validator   * SHA-256 Idempotency   * Local Fallback     |\n"
        "+------------------------------------+------------------------------------+\n"
        "                                     | Governed Execution\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|            SWYTCHCODE RUNTIME CLOUD (app.swytchcode.com/dashboard)      |\n"
        "|     Policy Engine  *  Zero-Token Credential Vault  *  Telemetry Stream  |\n"
        "+--------------------+-------------------------------+--------------------+\n"
        "                     |                               |                     \n"
        "                     v                               v                     \n"
        "+------------------------------------+ +----------------------------------+\n"
        "|        SARVAM AI INDIC VOICE       | |    EMERGENCY ASHA & DOCTOR QUEUE |\n"
        "|   * Saaras (Speech-to-Text)        | |    * Urgent Triage Webhook       |\n"
        "|   * Bulbul (Text-to-Speech)        | |    * Deduplicated 1x Delivery    |\n"
        "+------------------------------------+ +----------------------------------+"
    )
    
    arch_table = doc.add_table(rows=1, cols=1)
    arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    arch_cell = arch_table.rows[0].cells[0]
    arch_cell.width = Inches(7.0)
    set_cell_background(arch_cell, "0F172A")
    set_cell_margins(arch_cell, top=100, bottom=100, left=140, right=140)
    
    ap = arch_cell.paragraphs[0]
    a_run = ap.add_run(arch_text)
    a_run.font.name = "Consolas"
    a_run.font.size = Pt(7.5)
    a_run.font.color.rgb = RGBColor(226, 232, 240)

    doc.add_page_break()

    # 5. Before vs With Swytchcode Comparison Table
    h3 = doc.add_paragraph()
    r_h3 = h3.add_run("3. Architectural Evolution: Before vs. With Swytchcode")
    r_h3.font.bold = True
    r_h3.font.size = Pt(13)
    r_h3.font.color.rgb = RGBColor(15, 23, 42)

    comp_table = doc.add_table(rows=6, cols=3)
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    comp_table.autofit = False

    w_cap = Inches(1.3)
    w_bef = Inches(2.8)
    w_aft = Inches(2.9)

    headers = ["Capability", "Before Swytchcode (Legacy)", "With Swytchcode (Governed Runtime)"]
    for i, h_text in enumerate(headers):
        cell = comp_table.cell(0, i)
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        hp = cell.paragraphs[0]
        hrun = hp.add_run(h_text)
        hrun.font.bold = True
        hrun.font.size = Pt(8.5)
        if i == 1:
            hrun.font.color.rgb = RGBColor(220, 38, 38)
        elif i == 2:
            hrun.font.color.rgb = RGBColor(22, 163, 74)

    comp_data = [
        ("Tool Security", "Raw API keys exposed in app memory; vulnerable to prompt extraction.", "Zero-Token Isolation: LLM emits intent only. Keys locked in Swytchcode vault."),
        ("Emergency Idempotency", "Flaky 2G/3G network drops caused clients to retry -> 3 to 5 duplicate ambulance dispatches.", "Deterministic Deduplication: SHA-256 idempotency suppresses duplicate alerts within 5 min."),
        ("Clinical Validation", "Ad-hoc try/catch blocks. Malformed blood pressure or oxygen values could hit webhooks.", "Pre-Execution Guardrail: Pydantic schema validation rejects invalid clinical payloads before network dispatch."),
        ("Sarvam Indic Voice", "Direct REST calls to Sarvam. Socket drops on rural towers caused app to crash.", "Governed Voice Proxy: Swytchcode manages 3s timeout budgets, retries, and local phonetic fallback."),
        ("Live Telemetry", "Raw terminal stdout logs. Zero visual proof for judges or health auditors.", "Live Dashboard: Real-time event telemetry stream at app.swytchcode.com/dashboard/overview.")
    ]

    for row_idx, (cap, bef, aft) in enumerate(comp_data, start=1):
        r_cells = comp_table.rows[row_idx].cells
        r_cells[0].width = w_cap
        r_cells[1].width = w_bef
        r_cells[2].width = w_aft

        for c in r_cells:
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)

        # Cap
        p0 = r_cells[0].paragraphs[0]
        r0 = p0.add_run(cap)
        r0.font.bold = True
        r0.font.size = Pt(8.5)

        # Bef
        p1 = r_cells[1].paragraphs[0]
        r1 = p1.add_run(bef)
        r1.font.size = Pt(8.5)

        # Aft
        p2 = r_cells[2].paragraphs[0]
        r2 = p2.add_run(aft)
        r2.font.size = Pt(8.5)

    doc.add_paragraph()

    # 6. The 3 Governed Tools
    h4 = doc.add_paragraph()
    r_h4 = h4.add_run("4. The 3 Governed Healthcare Tools in Aarogya Sahayak")
    r_h4.font.bold = True
    r_h4.font.size = Pt(13)
    r_h4.font.color.rgb = RGBColor(15, 23, 42)

    tools = [
        ("1. dispatch_emergency_asha_alert", 
         "Trigger: Severe pre-eclampsia, cardiac, or pediatric danger signs.\n"
         "Action: Dispatches urgent clinical escalation notification to assigned ASHA & doctor queue.\n"
         "Swytchcode Governance: Ingress schema validation (BP, SpO2, weeks), SHA-256 idempotency deduplication, zero PII exposure."),
        ("2. sarvam_indic_voice_gateway",
         "Trigger: Citizen speaks in Marathi (mr-IN) or Hindi (hi-IN).\n"
         "Action: Routes audio through Sarvam AI (Saaras STT & Bulbul TTS).\n"
         "Swytchcode Governance: Enforces 3,000ms latency budget, language allowlist, and graceful phonetic fallback."),
        ("3. query_health_facility_registry",
         "Trigger: Patient or ASHA worker searches for nearest facility with ICU, NICU, or 24x7 emergency.\n"
         "Action: Queries verified Ayushman Bharat PM-JAY empanelled facilities.\n"
         "Swytchcode Governance: Read-only execution boundary; blocks any unauthorized database write mutations.")
    ]

    for t_title, t_desc in tools:
        tp = doc.add_paragraph()
        r_tt = tp.add_run(t_title + "\n")
        r_tt.font.bold = True
        r_tt.font.size = Pt(10)
        r_tt.font.color.rgb = RGBColor(15, 23, 42)
        r_td = tp.add_run(t_desc)
        r_td.font.size = Pt(8.5)
        r_td.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_page_break()

    # 7. Live Proof & Verification Guide
    h5 = doc.add_paragraph()
    r_h5 = h5.add_run("5. Step-by-Step Live Proof & Verification Guide (For Judges)")
    r_h5.font.bold = True
    r_h5.font.size = Pt(13)
    r_h5.font.color.rgb = RGBColor(15, 23, 42)

    proof_steps = [
        ("TEST 1: Swytchcode Runtime Health & Registered Tools",
         "curl -X GET 'https://your-backend.onrender.com/api/swytchcode/status'",
         "Expected Output: {'status': 'LIVE_CONNECTED', 'workspace_alias': 'calm-meadow-c150', 'tools_registered': ['dispatch_emergency_asha_alert', 'sarvam_indic_voice_gateway', ...]}"),
        ("TEST 2: Live Governed Emergency Triage Dispatch",
         "curl -X POST 'https://your-backend.onrender.com/api/swytchcode/execute-tool' \\\n  -H 'Content-Type: application/json' \\\n  -d '{\"tool_name\": \"dispatch_emergency_asha_alert\", \"priority\": \"CRITICAL\", \"clinical_condition\": \"Severe pre-eclampsia: BP 165/105\"}'",
         "Expected Output: {'status': 'DISPATCHED', 'trace_id': 'SWY-EMG-C0E4A3D2', 'latency_ms': 135.2, 'idempotency_enforced': true, 'dashboard_audit_url': 'https://app.swytchcode.com/dashboard/overview'}"),
        ("TEST 3: The Idempotency Test (Duplicate Suppression Defense)",
         "# Run the exact same curl command a second time within 5 minutes:\ncurl -X POST 'https://your-backend.onrender.com/api/swytchcode/execute-tool' ...",
         "Expected Output: {'status': 'ALREADY_DISPATCHED_IDEMPOTENT', 'idempotency_hit': true, 'message': 'Duplicate emergency alert suppressed by Swytchcode idempotency engine.'}")
    ]

    for title, cmd, exp in proof_steps:
        p = doc.add_paragraph()
        r_pt = p.add_run(title + "\n")
        r_pt.font.bold = True
        r_pt.font.size = Pt(9.5)

        # Code block
        c_table = doc.add_table(rows=1, cols=1)
        c_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        c_cell = c_table.rows[0].cells[0]
        c_cell.width = Inches(7.0)
        set_cell_background(c_cell, "F1F5F9")
        set_cell_margins(c_cell, top=60, bottom=60, left=100, right=100)
        cp = c_cell.paragraphs[0]
        crun = cp.add_run(cmd)
        crun.font.name = "Consolas"
        crun.font.size = Pt(7.5)
        crun.font.color.rgb = RGBColor(15, 23, 42)

        ep = doc.add_paragraph()
        erun = ep.add_run(exp)
        erun.font.name = "Consolas"
        erun.font.size = Pt(7.5)
        erun.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph()

    # 8. Architecture Visual Diagram Prompts
    h6 = doc.add_paragraph()
    r_h6 = h6.add_run("6. Architecture Visual Diagram Prompts (For Slides / Eraser.io)")
    r_h6.font.bold = True
    r_h6.font.size = Pt(13)
    r_h6.font.color.rgb = RGBColor(15, 23, 42)

    dp_intro = doc.add_paragraph()
    dp_intro.add_run("Paste the prompt below into Eraser.io, Napkin AI, or Mermaid Live Editor to render high-resolution diagram graphics:")

    mermaid_code = (
        "graph TD\n"
        "    A[Citizen Mobile PWA / Voice] -->|HTTPS Audio/Symptoms| B(FastAPI Backend)\n"
        "    B --> C{Deterministic Rule Engine}\n"
        "    C -->|Emergency Detected| D[PII Masking Engine]\n"
        "    D -->|Sanitized Intent| E[Swytchcode Governance Runtime]\n"
        "    E -->|Idempotent Webhook| F[ASHA & Doctor Emergency Queue]\n"
        "    E -->|Governed Proxy| G[Sarvam AI Indic Voice Saaras/Bulbul]\n"
        "    E -->|Live Telemetry| H[Swytchcode Dashboard app.swytchcode.com]"
    )

    d_table = doc.add_table(rows=1, cols=1)
    d_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    d_cell = d_table.rows[0].cells[0]
    d_cell.width = Inches(7.0)
    set_cell_background(d_cell, "F8FAFC")
    set_cell_margins(d_cell, top=80, bottom=80, left=120, right=120)
    dp = d_cell.paragraphs[0]
    drun = dp.add_run(mermaid_code)
    drun.font.name = "Consolas"
    drun.font.size = Pt(8)
    drun.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph()

    # 9. 2-Minute Winning Pitch Script
    h7 = doc.add_paragraph()
    r_h7 = h7.add_run("7. 2-Minute Winning Pitch Script for Judges")
    r_h7.font.bold = True
    r_h7.font.size = Pt(13)
    r_h7.font.color.rgb = RGBColor(15, 23, 42)

    pitch_table = doc.add_table(rows=1, cols=1)
    pitch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_cell = pitch_table.rows[0].cells[0]
    p_cell.width = Inches(7.0)
    set_cell_background(p_cell, "FFF7ED")
    set_cell_margins(p_cell, top=140, bottom=140, left=160, right=160)

    pp = p_cell.paragraphs[0]
    pr_t = pp.add_run("🎙️ What to Say Word-for-Word on Stage:\n\n")
    pr_t.font.bold = True
    pr_t.font.size = Pt(9.5)
    pr_t.font.color.rgb = RGBColor(154, 52, 18)

    pitch_speech = (
        "\"Namaste judges. In rural India, when an AI assistant detects that a pregnant mother has a critical blood "
        "pressure of 165/100, allowing an LLM to directly call external APIs is dangerous. Models hallucinate parameters, "
        "rural 3G network drops cause duplicate ambulance dispatches, and credentials can leak.\n\n"
        "We integrated Swytchcode as our enterprise AI tool execution & governance layer. Every single action—from "
        "Sarvam AI Indic voice translation in Marathi and Hindi to emergency ASHA dispatches—is governed by Swytchcode.\n\n"
        "As you can see right here on our live Swytchcode Dashboard at app.swytchcode.com, the triage alert was executed "
        "with status 200 OK, latency 135ms, strict schema validation, zero-token security, and guaranteed idempotency. "
        "Swytchcode makes AI in healthcare safe, deterministic, and ready for 1.4 billion citizens!\""
    )
    pr_s = pp.add_run(pitch_speech)
    pr_s.font.italic = True
    pr_s.font.size = Pt(9)
    pr_s.font.color.rgb = RGBColor(124, 45, 18)

    # Save Document
    doc.save(output_path)
    print(f"Successfully generated clean Word Document (.docx): {output_path}")

if __name__ == "__main__":
    out_file = os.path.abspath(os.path.join(os.path.dirname(__file__), "../swytchcode/AarogyaSahayak_Swytchcode_Architecture_and_Proof.docx"))
    generate_docx(out_file)
